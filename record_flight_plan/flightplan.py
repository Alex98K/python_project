from itertools import groupby
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import doc_to_docx
import docx
import time
import xlrd
import os

CN_NUM = {1: ('一', '（一）'), 2: ('二', '（二）'), 3: ('三', '（三）'), 4: ('四', '（四）'), 5: ('五', '（五）'),
          6: ('六', '（六）'), 7: ('七', '（七）'), 8: ('八', '（八）'), 9: ('九', '（九）'), 10: ('十', '（十）'),
          11: ('十一', '（十一）'), 12: ('十二', '（十二）'), 13: ('十三', '（十三）'), 14: ('十四', '（十四）'),
          15: ('十五', '（十五）'), 16: ('十六', '（十六）'), 17: ('十七', '（十七）'), 18: ('十八', '（十八）'),
          19: ('十九', '（十九）'), 20: ('二十', '（二十）'), '正班': '增补正班计划', '加班': '加班航班计划', '临时经营': '临时经营航班计划',
          '北京': 1, '大兴': 2, '天津': 3, '河北': 4, '山西': 5, '内蒙古': 6,
          '东北': 1, '华东': 2, '中南': 3, '西南': 4, '西北': 5, '新疆': 6, '包机': '临时经营航班计划'
          }
HANG_JI = '2020夏秋'  # 航季，主要用在正文中什么什么航季地方
TABLE_SAMPLE = 'chaosong_table'  # 在word模板中建立的自定义表格样式名字,主要是有上中下边框，没其他边框
PATH = os.path.abspath(os.path.dirname(__file__))


def eng2chin(string, str_type):
    """
    # 把公司二字码和机场四字码转化为中文名称,并且返回机场和航空公司的抄送单位，管理局和监管局
    :param string: 输入字符串或列表，机场四字码或者公司二字码列表
    :param str_type:指出是机场还是公司
    :return:返回中文，抄送管理局、监管局
    """
    if str_type == 'airline':
        airline_name = airline_record[string][col_name_num_airline['航空公司中文名']]
        cc_company = airport_record_3w[airline_record[string][col_name_num_airline['基地机场三字码']]][
            col_name_num_airport['管理局']]
        if cc_company == '华北':
            cc_company = ''
        return airline_name, cc_company
    elif str_type == 'airport':
        airport_name = []  # 初始化机场航线列表
        cc_airport_jian = []  # 初始化抄送监管局列表
        cc_airport_guan = []  # 初始化抄送外地管理局列表
        for j1 in string.split('-'):
            airport_name.append(airport_record_4w[j1][col_name_num_airport['机场中文名']])
            cc_pro = airport_record_4w[j1][col_name_num_airport['省份中文名']]
            cc_ju = airport_record_4w[j1][col_name_num_airport['管理局']]
            # 判断是管理局是华北的，就是抄送监管局，不是就抄送管理局；监管局分为北京、大兴和其他监管局
            if cc_ju == '华北':
                if j1 == 'ZBAD':
                    cc_airport_jian.append('大兴')
                elif j1 == 'ZBAA':
                    cc_airport_jian.append('北京')
                elif j1 == 'ZBNY':
                    raise Exception('竟然还有南苑机场')
                else:
                    cc_airport_jian.append(cc_pro)
            else:
                cc_airport_guan.append(cc_ju)
        airport_name = '-'.join(airport_name)  # 列表连接成字符串
        cc_airport_guan = list(set(cc_airport_guan))  # 去重
        cc_airport_jian = list(set(cc_airport_jian))  # 去重
        return airport_name, cc_airport_guan, cc_airport_jian


def cc_trans(cc_guan_t):
    """
    # 用来将输入的机场或航空公司列表转为可以输出的字符串,如果输入的是['ha', ['hah','haha']]就先判断每个元素的类型
    """
    list_zu = []
    for x1 in cc_guan_t:
        if isinstance(x1, list):
            for y in x1:
                list_zu.append(y)
        elif isinstance(x1, str) and x1 != '':
            list_zu.append(x1)
        elif not x1:
            pass
    temp1 = sorted(list(set(list_zu)), key=lambda k1: CN_NUM[k1])
    return '、'.join(temp1)


def text_font(paragragh, text, font_pos='正文'):
    """
    # 正文字体函数，定义三个样式函数，字体、段落、表格
    :param text: 输入的文字内容，字符串
    :param paragragh: 段落
    :param font_pos: 字体类型，默认仿宋
    """
    run = paragragh.add_run(text)
    run.font.size = Pt(16)  # 默认三号字体，正文和抄送使用，如果改变用途为标题，再改大小
    if font_pos == '标题':
        run.font.size = Pt(22)  # 标题的字体
        font = u'方正小标宋_GBK'
    elif font_pos == '抄送':
        font = u'楷体_GB2312'
    elif font_pos == '电报头':
        run.font.size = Pt(70)  # 标题的字体
        font = u'方正大标宋简体'
    elif font_pos == '加急':
        font = u'黑体'
    elif font_pos == '文件头空白':
        font = u'黑体'
    else:
        font = u'仿宋_GB2312'
    run.font.name = font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)  # 中文字体特殊设置方法


def para_style_input(doc, text='', font_pos='正文'):
    """
    # 正文段落格式，默认两端对齐，定义三个样式函数，字体、段落、表格
    返回段落指引，可后续添加文本
    """
    paragraph = doc.add_paragraph()
    if font_pos == '发往单位':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif font_pos == '正文落款':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif font_pos == '标题':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = Pt(30)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    elif font_pos == '电报头':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
    else:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if font_pos != '签批':
            paragraph.paragraph_format.first_line_indent = Inches(0.45)  # 首行缩进
    if text != '':
        text_font(paragraph, text, font_pos)
    return paragraph


def table_para_style_input(table_cell, text, style='居中', font_pos='正文'):
    """
    # 表格格式及输入文字函数，定义三个样式函数，字体、段落、表格
    """
    para_table1 = table_cell.paragraphs[0]  # 表格中的特殊用法
    if style == '首行缩进':  # 表格中的首行缩进文字，例如航线，日期那行
        para_table1.paragraph_format.first_line_indent = Inches(0.45)
    elif style == '居中':  # 表格中的居中文字，例如班期，航班号那些格子
        para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style == '悬垂缩进':  # 抄送表格中的字体，注意下面的用法，悬垂缩进
        para_table1.paragraph_format.left_indent = Inches(0.55)  # 首先左对齐
        para_table1.paragraph_format.first_line_indent = Inches(-0.65)  # 然后用负数表示悬垂缩进
    elif style == '两端对齐':
        para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text_font(para_table1, text, font_pos)  # 表格中首行缩进、居中文字输入，并设置字体
    return para_table1


def title_style_new_input(doc, text):
    # 以下内容为使用管理局电报模板后，从标题那里输入时的标题内容及格式
    para = doc.paragraphs[-1]
    run = para.add_run(text)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(30)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run.font.size = Pt(22)  # 标题的字体
    font = u'方正小标宋_GBK'
    run.font.name = font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)  # 中文字体特殊设置方法


# 开始计时
time_start = time.time()

# 读取航班计划备案excel文件
book = xlrd.open_workbook(os.path.join(PATH, '国内航班计划备案.xlsx'))
data = book.sheet_by_index(0)

# 生成表头和列数的对应关系，字典类型，并打印
col_name_num = {key: value for value, key in enumerate(data.row_values(1))}
# print(col_name_num)

# 生成航班计划记录，并且将读取出来是float类型的转化为str
flight_record = [(list(map(lambda arg: str(arg).replace('.0', ''), data.row_values(row))))
                 for row in range(2, data.nrows)]
# 判断记录中是否有重复值
for u in flight_record:
    if flight_record.count(u) > 1:
        raise ('excel中存在重复的数据\n{}'.format(u))
# print(flight_record)

flight_plan = []  # 初始化输出最终结果航班计划列表
flight_record.sort(key=lambda p4: p4[col_name_num['航班号']])  # 将航班记录按照航班号排序
for j in groupby(flight_record, key=lambda p1: p1[col_name_num['航司二字码']]):  # 在groupby按照公司二字码进行分组后，j 代表每一个公司
    # print('\n\n***************下一个公司****************\n\n', 'j[1]', j[1])
    f_company_job = sorted(j[1], key=lambda p2: p2[col_name_num['任务性质']])  # 将每个公司的航班计划按照任务性质排序
    # print('f_company_job', f_company_job)
    job_plan = []  # 初始化按任务性质分类的最终航班计划
    for p in groupby(f_company_job, key=lambda p3: p3[col_name_num['任务性质']]):  # 在groupby按照任务性质进行分组后，p 代表每一个任务性质：加班、正班
        # print('\n***********-----%s-----***********\n\n' % p[0], 'j[0]', j[0], 'p[0]', p[0], 'p[1]', p[1])
        jihua_temp = []  # 初始化每一个公司每个任务类型的航班计划最终列表
        for x in p[1]:  # x 表示每一条航班计划
            # print('下一个航班计划----\nx {}\njihua_temp  {}'.format(x, jihua_temp))
            def fun(shuzu):
                shuzu_fl_nb = shuzu[col_name_num['航班号']].replace(j[0], '')  # 有可能是新加入的航班计划
                x_fl_nb = x[col_name_num['航班号']].replace(j[0], '')  # 之前列表中的航班计划
                # 判断后一个航班计划与之前存在列表中的是否为同一航班，如果后面的航班号是前面的加减1，其他相同，那就是一个航班
                if (shuzu[col_name_num['班期']] == x[col_name_num['班期']]) and \
                        (shuzu[col_name_num['生效日期']] == x[col_name_num['生效日期']]) and \
                        (shuzu[col_name_num['截止日期']] == x[col_name_num['截止日期']]) and \
                        ('-'.join(list(reversed(shuzu[col_name_num['航线四字码']].split('-')))) ==
                         x[col_name_num['航线四字码']]) and (abs(int(shuzu_fl_nb) - int(x_fl_nb)) == 1):
                    # 进行航班计划变更，把原来列表中的航班计划变更为最后的航班号，航线，例如CA1234/5,AB-YY-AB
                    if int(shuzu_fl_nb) % 2 == 1:
                        # 航班号是奇数的，就奇数的在前偶数在后
                        fl_after = [v for i, v in enumerate(x_fl_nb) if shuzu_fl_nb[i] != v]
                        shuzu[col_name_num['航班号']] = j[0] + shuzu_fl_nb + '/' + "".join(fl_after)
                        shuzu[col_name_num['航线四字码']] = \
                            shuzu[col_name_num['航线四字码']] + '-' + x[col_name_num['航线四字码']].split('-', 1)[1]
                    else:
                        fl_after = [v for i, v in enumerate(shuzu_fl_nb) if x_fl_nb[i] != v]
                        shuzu[col_name_num['航班号']] = j[0] + x_fl_nb + '/' + "".join(fl_after)
                        shuzu[col_name_num['航线四字码']] = \
                            x[col_name_num['航线四字码']] + '-' + shuzu[col_name_num['航线四字码']].split('-', 1)[1]
                    # print('fun inner', shuzu)
                    return True
                else:
                    return False


            if not jihua_temp:
                jihua_temp.append(x)
            else:
                if list(filter(fun, jihua_temp)):  # 用filter函数对之前保存的航班计划列表进行筛选，剩下的就是与计划加入航班计划为同一个航班的
                    # print('过滤if 正确的', jihua_temp)
                    pass
                else:
                    jihua_temp.append(x)  # 不是同一航班的就加入航班计划列表中
        # print("jihua_temp:")
        # pprint(jihua_temp, compact=True)
        job_plan.append([p[0], jihua_temp])
        # print('job_plan', job_plan)
    # flight_plan 的结构如下
    # [[ca, [[正班, [[ca,ca1234],[ca,ca4321]]], [加班, [jihua1,jihua2]]]],\
    # [sc, [[正班, [jihua1,jihua2]], [加班, [jihua1,jihua2]]]],\
    # [9c, [[正班, [jihua1,jihua2]], [加班, [jihua1,jihua2]]]]]
    flight_plan.append([j[0], job_plan])
# 打印最终形成的航班计划列表
# print('***********最终形成的航班计划列表是***********')
# pprint(flight_plan, compact=True)

# 导入机场和航空公司、受限信息信息表
book = xlrd.open_workbook(os.path.join(PATH, 'flightplan/info.xlsx'))
data_airport = book.sheet_by_name('airport')
data_airline = book.sheet_by_name('airline')
data_restrict = book.sheet_by_name('restrict')

# 生成列名、列数的字典，备查,并打印
col_name_num_airport = {key: value for value, key in enumerate(data_airport.row_values(0))}
col_name_num_airline = {key: value for value, key in enumerate(data_airline.row_values(0))}
col_name_num_restrict = {key: value for value, key in enumerate(data_restrict.row_values(0))}
# pprint(col_name_num_airport)
# pprint(col_name_num_airline)
# print(col_name_num_restrict)

# 分别以机场4字码和3字码、公司2字码、受限信息，生成信息字典，可快速查询信息，并打印
airport_record_4w = {data_airport.row_values(row)[1]: data_airport.row_values(row)
                     for row in range(1, data_airport.nrows)}
airport_record_3w = {data_airport.row_values(row)[0]: data_airport.row_values(row)
                     for row in range(1, data_airport.nrows)}
airline_record = {data_airline.row_values(row)[0]: data_airline.row_values(row) for row in range(1, data_airline.nrows)}
restrict_record = {data_restrict.row_values(0)[row]: data_restrict.col_values(row)[1:]
                   for row in range(0, data_restrict.ncols)}

# pprint(airport_record_4w, compact=True)
# pprint(airport_record_3w, compact=True)
# pprint(airline_record, compact=True)
# print(restrict_record)


# 新建word文档，并输入文件开头几段
# document = docx.Document(os.path.join(PATH, 'flightplan\\telegraph_template.docx'))
document = docx.Document(os.path.join(PATH, 'flightplan\\template.docx'))

# 输入文件开头,套用模板,以下内容因为年代那里涉及时间域等设置，与管理局电报模板不同，因此改用套管理局电报模板方式，以下内容暂时不使用
# # para_style_input(document, font_pos='文件头空白')
# para_style_input(document, '民航明传电报', font_pos='电报头')
# para = para_style_input(document, '发电单位', font_pos='签批')
# para.paragraph_format.space_before = Pt(18)
# text_font(para, ' 民航华北地区管理局', font_pos='抄送')
# text_font(para, '               签发盖章')
# table = document.add_table(rows=1, cols=2, style=TABLE_SAMPLE)
# para1 = table_para_style_input(table.cell(0, 0), '等级 ', style='两端对齐')
# text_font(para1, '加急·明电', font_pos='加急')
# table_para_style_input(table.cell(0, 1), '华北局发明电〔{}〕    号'
#                        .format(time.localtime(time.time()).tm_year), style='两端对齐', font_pos='抄送') \
#     .paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 输入文件标题
xuhao_len = len(flight_plan)  # 获取有多少家公司--即大序号一、二、三
if xuhao_len == 1:
    title = '关于同意{}部分航班计划备案事项的报告'.format(airline_record[flight_plan[0][0]][col_name_num_airline['中文简称']])
else:
    title = '关于同意{}等公司部分航班计划备案事项的报告'.format(airline_record[flight_plan[0][0]][col_name_num_airline['中文简称']])
# para_style_input(document, title, font_pos='标题')  # 改用套管理局电报模板方式后，暂时不使用
title_style_new_input(document, title)

# 输入发电单位和正文一部分
para_style_input(document, '民航局运行监控中心：', font_pos='发往单位')
para_style_input(document, '根据《中国民航国内航线航班评审规则》要求，经审核，我局同意以下航空公司{}航季期间备案部分航班计划，具体如下：'.format(HANG_JI))

cc_guan = []  # 初始化抄送管理局列表
cc_jian = []  # 初始化抄送监管局列表
cc_comp = []  # 初始化抄送航空公司列表
cc_airport = []  # 初始化抄送机场列表
for i_comp, v_comp in enumerate(flight_plan):  # 获取各个公司,i_comp是公司序号
    company_name, cc_guan_temp = eng2chin(v_comp[0], 'airline')  # 获取公司全称和抄送管理局
    cc_guan.append(cc_guan_temp)  # 添加到抄送管理局
    company_name_simple = airline_record[v_comp[0]][col_name_num_airline['中文简称']]
    cc_comp.append(company_name_simple)  # 添加到抄送公司
    if xuhao_len == 1:  # 如果只有一家公司，就不需要大的公司前面的序号
        para3 = para_style_input(document, '{company_name}'.format(company_name=company_name))
    else:
        para3 = para_style_input(document, '{xuhao}、{company_name}'.format(xuhao=CN_NUM[i_comp + 1][0],
                                                                           company_name=company_name), font_pos='加急')
    job_len = len(v_comp[1])  # 获取有几种任务性质--即小序号（一）（二）（三）
    for i_job, v_job in enumerate(v_comp[1]):  # 获取每家公司的任务类型，i_job是任务序号
        job_name = v_job[0]  # 任务类型名称
        if job_len == 1:  # 如果只有一种任务性质，就不需要任务性质前面的小序号
            if xuhao_len == 1:
                text_font(para3, '{job_name}'.format(job_name=CN_NUM[job_name]))  # 只有一种任务性质时，追加到公司名称后面
            else:
                text_font(para3, '{job_name}'.format(job_name=CN_NUM[job_name]), font_pos='加急')  # 只有一种任务性质时，追加到公司名称后面
        else:
            # 任务性质那行输出
            para_style_input(document, '{xuhao}{job_name}'.format(xuhao=CN_NUM[i_job + 1][1],
                                                                  job_name=CN_NUM[job_name]), font_pos='抄送')
        hang_ban_num = len(v_job[1])  # 如果航班计划条数大于1，就在航班号前面加上序号
        for index, value in enumerate(v_job[1]):  # 准备输入每一个航班计划
            # 把航线变成中文，并获取机场抄送管理局、监管局名单
            hangxian, cc_guan_temp_value, cc_jian_temp = eng2chin(value[col_name_num['航线四字码']], 'airport')
            cc_airport.append(hangxian)  # 添加到抄送机场
            cc_guan.append(cc_guan_temp_value)  # 添加到抄送管理局
            cc_jian.append(cc_jian_temp)  # 添加到抄送监管局
            same_date = value[col_name_num['生效日期']] == value[col_name_num['截止日期']]
            table = document.add_table(rows=4, cols=3, style='Normal Table')  # 准备写入航班计划，新建表格
            if hang_ban_num > 1:  # 如果航班计划条数大于1，就在航班号前面加上序号
                table_para_style_input(table.cell(0, 0), f'{index+1}、航班号')
            else:
                table_para_style_input(table.cell(0, 0), '航班号')
            table_para_style_input(table.cell(0, 1), '机型')
            if not same_date:
                table_para_style_input(table.cell(0, 2), '班期')
            table_para_style_input(table.cell(1, 0), value[col_name_num['航班号']])
            table_para_style_input(table.cell(1, 1), '不限')
            if not same_date:
                table_para_style_input(table.cell(1, 2), value[col_name_num['班期']].replace('.', ''))
            cell1 = table.cell(2, 0)
            cell2 = table.cell(2, 2)
            cell1.merge(cell2)  # 合并单元格，输入航线
            table_para_style_input(cell1, hangxian, style='首行缩进')
            cell1 = table.cell(3, 0)
            cell2 = table.cell(3, 2)
            cell1.merge(cell2)  # 合并单元格，输入日期
            if same_date:  # 执行日期只有一天的，不显示截止日期，更改格式，去掉03月05日中的0
                date_time = time.strptime(value[col_name_num['生效日期']], '%Y-%m-%d')
                table_para_style_input(cell1, '航班执行日期：' + '{}年{}月{}日'.format(date_time.tm_year, date_time.tm_mon,
                                                                             date_time.tm_mday), style='首行缩进')
            else:  # 执行日期不是一天的，显示日期，更改格式，去掉03月05日中的0
                date_time_start = time.strptime(value[col_name_num['生效日期']], '%Y-%m-%d')
                date_time_end = time.strptime(value[col_name_num['截止日期']], '%Y-%m-%d')
                table_para_style_input(cell1, '航班执行日期：' + '{}年{}月{}日'
                                       .format(date_time_start.tm_year, date_time_start.tm_mon, date_time_start.tm_mday)
                                       + '-' + '{}年{}月{}日'.format(date_time_end.tm_year, date_time_end.tm_mon,
                                                                  date_time_end.tm_mday), style='首行缩进')
            # 如果公司或机场受限制，采用置换方式，就在航班计划表格之后加一个括号，采用置换的,取消备注中的内容
            if company_name_simple in restrict_record['受限公司'] \
                    or ([p7 for p7 in hangxian.split('-') if p7 in restrict_record['受限机场']]):
                para_style_input(document, '（置换：{}）'.format(value[col_name_num['备注']]))
# 将抄送的单位转为可以输出的字符串
cc_guan = cc_trans(cc_guan)
cc_jian = cc_trans(cc_jian)
cc_comp = '、'.join(cc_comp)
temp = []
for k in cc_airport:
    temp.extend(k.split('-'))
cc_airport = '、'.join(list(set(temp)))

# 形成正文尾部
para_style_input(document, '此报告。')
para_style_input(document, '')
para_style_input(document, '民航华北地区管理局', font_pos='正文落款')
nowadays = time.localtime(time.time())
para_style_input(document, '{}'.format(time.strftime("{}年{}月{}日".format(
    nowadays.tm_year, nowadays.tm_mon, nowadays.tm_mday))), font_pos='正文落款')

# 文件末尾的抄送表格
table = document.add_table(rows=2, cols=1, style=TABLE_SAMPLE)
table_para_style_input(table.cell(0, 0), '抄送：', style='悬垂缩进', font_pos='抄送')
para2 = table_para_style_input(table.cell(0, 0), '民航局运输司；', style='悬垂缩进')
if cc_guan:
    text_font(para2, '{}地区管理局，'.format(cc_guan))
if cc_jian:
    text_font(para2, '{}监管局，'.format(cc_jian))
text_font(para2, '{}，{}机场，华北空管局，华北局京津冀办。'.format(cc_comp, cc_airport))
table_para_style_input(table.cell(0, 1), '承办单位：', style='两端对齐', font_pos='抄送')
table_para_style_input(table.cell(0, 1), '华北局运输管理处          ', style='两端对齐')
table_para_style_input(table.cell(0, 1), '电话：', style='两端对齐', font_pos='抄送')
table_para_style_input(table.cell(0, 1), '010-64593776', style='两端对齐')

# 保存文件
document.save(os.path.join(PATH, '{}.docx'.format(title)))
print('\n*********docx文件成功生成！用时{}秒，准备将docx文件转为doc文件......*********\n'.format((time_now := time.time()) - time_start))
doc_to_docx.docx2doc(title)
print('*********转doc文件成功！用时{}秒*********'.format((time.time()) - time_now))
