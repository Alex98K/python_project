import re
import docx
import pymysql
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt


class RawData(object):
    def __init__(self, start_date, end_date):
        self.db = pymysql.connect(host="localhost", user="root", passwd="jiajia0611", db="ceshi")
        self.cur = self.db.cursor()
        self.procedure = 'p2'
        pattern = r"^20[0-9][0-9][0123][0-9][0123][0-9]$"
        if not (re.search(pattern, start_time) and re.search(pattern, end_time)):
            raise Exception('请输入正确的日期，例如：20200101')
        self.start_time = start_date
        self.end_time = end_date

    def __del__(self):
        self.cur.close()
        self.db.close()

    @property
    def data_all(self):
        """
        获取通报中所需要的全部原始数据
        """
        try:
            self.cur.callproc(self.procedure, (self.start_time, self.end_time, 0))
            res = [self.cur.fetchall()]
            title = [tuple(j[0] for j in self.cur.description)]
            while self.cur.nextset():
                if temp := self.cur.fetchall():
                    res.append(temp)
                    title.append(tuple(j[0] for j in self.cur.description))
        except Exception as e:
            print("调用存储过程失败")
            print("出错的原因是：", e)
            res = None
            title = None
        return res, title

    def data_one(self, logo):
        """
        获取一条原始数据
        """
        try:
            self.cur.callproc(self.procedure, (self.start_time, self.end_time, int(logo)))
            res = [self.cur.fetchall()]
            while self.cur.nextset():
                res.append(self.cur.fetchall())
        except Exception as e:
            print("调用存储过程失败")
            print("出错的原因是：", e)
            res = None
        return res


class WordWrite(object):
    def __init__(self):
        # 新建word文档，并输入文件开头几段
        self.document = docx.Document('flightplan\\telegraph_template.docx')

    def text_font(self, paragragh, text, font_pos='正文'):
        """
        # 正文字体函数，定义三个样式函数，字体、段落、表格
        :param text: 输入的文字内容，字符串
        :param paragragh: 段落
        :param font_pos: 字体类型，默认仿宋
        """
        run = paragragh.add_run(text)
        run.font.size = Pt(16)  # 默认三号字体，正文和抄送使用，如果改变用途为标题，再改大小
        if font_pos == '标题':
            run.font.size = Pt(22)  # 标题的字体
            font = u'方正小标宋_GBK'
        elif font_pos == '抄送':
            font = u'楷体_GB2312'
        elif font_pos == '电报头':
            run.font.size = Pt(70)  # 标题的字体
            font = u'方正大标宋简体'
        elif font_pos == '加急':
            font = u'黑体'
        elif font_pos == '文件头空白':
            font = u'黑体'
        else:
            font = u'仿宋_GB2312'
        run.font.name = font
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font)  # 中文字体特殊设置方法


    def para_style_input(self, doc, text='', font_pos='正文'):
        """
        # 正文段落格式，默认两端对齐，定义三个样式函数，字体、段落、表格
        返回段落指引，可后续添加文本
        """
        paragraph = doc.add_paragraph()
        if font_pos == '发往单位':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif font_pos == '正文落款':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif font_pos == '标题':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = Pt(30)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
        elif font_pos == '电报头':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(18)
        else:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if font_pos != '签批':
                paragraph.paragraph_format.first_line_indent = Inches(0.45)  # 首行缩进
        if text != '':
            self.text_font(paragraph, text, font_pos)
        return paragraph


    def table_para_style_input(self, table_cell, text, style='居中', font_pos='正文'):
        """
        # 表格格式及输入文字函数，定义三个样式函数，字体、段落、表格
        """
        para_table1 = table_cell.paragraphs[0]  # 表格中的特殊用法
        if style == '首行缩进':  # 表格中的首行缩进文字，例如航线，日期那行
            para_table1.paragraph_format.first_line_indent = Inches(0.45)
        elif style == '居中':  # 表格中的居中文字，例如班期，航班号那些格子
            para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == '悬垂缩进':  # 抄送表格中的字体，注意下面的用法，悬垂缩进
            para_table1.paragraph_format.left_indent = Inches(0.55)  # 首先左对齐
            para_table1.paragraph_format.first_line_indent = Inches(-0.65)  # 然后用负数表示悬垂缩进
        elif style == '两端对齐':
            para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.text_font(para_table1, text, font_pos)  # 表格中首行缩进、居中文字输入，并设置字体
        return para_table1

#
# # 输入文件开头,套用模板
# # para_style_input(document, font_pos='文件头空白')
# para_style_input(document, '民航明传电报', font_pos='电报头')
# para = para_style_input(document, '发电单位', font_pos='签批')
# para.paragraph_format.space_before = Pt(18)
# text_font(para, ' 民航华北地区管理局', font_pos='抄送')
# text_font(para, '               签发盖章')
# table = document.add_table(rows=1, cols=2, style=TABLE_SAMPLE)
# para1 = table_para_style_input(table.cell(0, 0), '等级 ', style='两端对齐')
# text_font(para1, '加急·明电', font_pos='加急')
# table_para_style_input(table.cell(0, 1), '华北局发明电〔{}〕    号'
#                        .format(time.localtime(time.time()).tm_year), style='两端对齐', font_pos='抄送')\
#     .paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#
# # 输入文件标题
# para_style_input(document, title, font_pos='标题')
#
# # 输入发电单位和正文一部分
# para_style_input(document, '民航局运行监控中心：', font_pos='发往单位')
# para_style_input(document, '根据《中国民航国内航线航班评审规则》要求，经审核，我局同意以下航空公司{}航季期间备案部分航班计划，具体如下：'.format(HANG_JI))


if __name__ == '__main__':
    start_time = '20200101'
    end_time = '20200131'
    p = time.time()
    raw_data = RawData(start_time, end_time).data_all

    print(raw_data[0], raw_data[1])
    print(len(raw_data[0]), len(raw_data[1]))
    print(time.time()-p)
    # print(raw_data.data_one(30))
