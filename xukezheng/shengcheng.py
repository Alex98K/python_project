from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import re
import docx
import time
import xlrd
import os

PATH_PROGRAM = os.path.abspath(os.path.dirname(__file__))
PATH = "E:\e\航线航班"  # 航线经营许可表格
PASS_DENG = 1667
PASS_ZHU = 803


def text_font(paragragh, text, font_pos='正文'):
    """
    # 正文字体函数，定义三个样式函数，字体、段落、表格
    :param text: 输入的文字内容，字符串
    :param paragragh: 段落
    :param font_pos: 字体类型，默认仿宋
    """
    run = paragragh.add_run(text)
    run.font.size = Pt(16)  # 默认三号字体，正文和抄送使用，如果改变用途为标题，再改大小
    if font_pos == '标题':
        run.font.size = Pt(22)  # 标题的字体
        font = u'方正小标宋_GBK'
    elif font_pos == '抄送':
        font = u'楷体_GB2312'
    elif font_pos == '电报头':
        run.font.size = Pt(70)  # 标题的字体
        font = u'方正大标宋简体'
    elif font_pos == '加急':
        font = u'黑体'
    elif font_pos == '文件头空白':
        font = u'黑体'
    else:
        font = u'仿宋_GB2312'
    run.font.name = font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)  # 中文字体特殊设置方法


def para_style_input(doc, text='', font_pos='正文'):
    """
    # 正文段落格式，默认两端对齐，定义三个样式函数，字体、段落、表格
    返回段落指引，可后续添加文本
    """
    paragraph = doc.add_paragraph()
    if font_pos == '发往单位':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif font_pos == '正文落款':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif font_pos == '标题':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = Pt(30)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    elif font_pos == '电报头':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
    else:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if font_pos != '签批':
            paragraph.paragraph_format.first_line_indent = Inches(0.45)  # 首行缩进
    if text != '':
        text_font(paragraph, text, font_pos)
    return paragraph


def table_para_style_input(table_cell, text, style='居中', font_pos='正文'):
    """
    # 表格格式及输入文字函数，定义三个样式函数，字体、段落、表格
    """
    para_table1 = table_cell.paragraphs[0]  # 表格中的特殊用法
    if style == '首行缩进':  # 表格中的首行缩进文字，例如航线，日期那行
        para_table1.paragraph_format.first_line_indent = Inches(0.45)
    elif style == '居中':  # 表格中的居中文字，例如班期，航班号那些格子
        para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif style == '悬垂缩进':  # 抄送表格中的字体，注意下面的用法，悬垂缩进
        para_table1.paragraph_format.left_indent = Inches(0.55)  # 首先左对齐
        para_table1.paragraph_format.first_line_indent = Inches(-0.65)  # 然后用负数表示悬垂缩进
    elif style == '两端对齐':
        para_table1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text_font(para_table1, text, font_pos)  # 表格中首行缩进、居中文字输入，并设置字体
    return para_table1


def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i1 in inline:
                if old_text in i1.text:
                    text = i1.text.replace(old_text, new_text)
                    i1.text = text


# 开始计时
time_start = time.time()

# 读取航班计划备案excel文件
# book = xlrd.open_workbook(os.path.join(PATH, '待做许可证的记录.xlsx'))
book = xlrd.open_workbook(os.path.join(PATH, '航线经营许可汇总表.xlsx'))
data_deng = book.sheet_by_name('登记')
data_zhu = book.sheet_by_name('注销')

ALL_deng = {}
for i in range(PASS_DENG, data_deng.nrows):
    temp = data_deng.row_values(i)
    try:
        temp[4] = int(temp[4])
    except ValueError:
        continue
    try:
        temp[5] = xlrd.xldate_as_tuple(temp[5], 0)
        temp[7] = xlrd.xldate_as_tuple(temp[7], 0)
    except TypeError:
        continue
    r = temp[1:6]
    r.extend(temp[7:])
    if ALL_deng.get(temp[6]):
        ALL_deng[temp[6]].append(r)
    else:
        ALL_deng[temp[6]] = [r]

for key, value in ALL_deng.items():
    # 新建word文档，并输入文件开头几段
    document = docx.Document(os.path.join(PATH_PROGRAM, '模板文件/登记.docx'))
    res = re.findall(r'〔(\d+)〕(\d+)', key)
    try:
        year = res[0][0]
        nb = res[0][1]
        company = value[0][1]
        com = value[0][0]
        month = value[0][5][1]
        day = value[0][5][2]
    except IndexError:
        continue
    replace_text(document, 'year', f'{year}')
    replace_text(document, 'nb', f'{nb}')
    replace_text(document, 'company', f'{company}')
    replace_text(document, 'month', f'{month}')
    replace_text(document, 'day', f'{day}')
    table = document.add_table(rows=len(value) + 1, cols=4, style="grid")  # 准备写入航班计划，新建表格
    table_para_style_input(table.cell(0, 0), '序号', font_pos='加急')
    table.cell(0, 0).width = Inches(0.2)
    table_para_style_input(table.cell(0, 1), '航线', font_pos='加急')
    table.cell(0, 1).width = Inches(7)
    table_para_style_input(table.cell(0, 2), '初始班次', font_pos='加急')
    table.cell(0, 2).width = Inches(0.5)
    table_para_style_input(table.cell(0, 3), '计划开航日期', font_pos='加急')
    table.cell(0, 3).width = Inches(3)
    for i, j in enumerate(value):
        air_line = j[2]
        hangbanliang = j[3]
        exe_date = j[4]
        table_para_style_input(table.cell(i + 1, 0), str(i + 1))
        table_para_style_input(table.cell(i + 1, 1), air_line)
        table_para_style_input(table.cell(i + 1, 2), str(hangbanliang))
        table_para_style_input(table.cell(i + 1, 3), '{}/{}/{}'.format(exe_date[0], exe_date[1], exe_date[2]))

    title = f'{year}{nb}-{com}-登记'
    path_deng = (os.path.join(PATH, '航线航班\航线经营许可\登记\{}'.format(year)))
    if not os.path.isdir(path_deng):
        os.mkdir(path_deng)
    path_deng_file = os.path.join(path_deng, '{}.docx'.format(title))
    if not os.path.isfile(path_deng_file):
        document.save(path_deng_file)
        print(path_deng_file)

ALL_zhu = {}
for i in range(PASS_ZHU, data_zhu.nrows):
    temp = data_zhu.row_values(i)
    try:
        temp[6] = xlrd.xldate_as_tuple(temp[6], 0)
    except TypeError:
        continue
    r = temp[1:4]
    r.extend(temp[6:])
    if ALL_zhu.get(temp[5]):
        ALL_zhu[temp[5]].append(r)
    else:
        ALL_zhu[temp[5]] = [r]

for key, value in ALL_zhu.items():
    # 新建word文档，并输入文件开头几段
    document = docx.Document(os.path.join(PATH_PROGRAM, '模板文件/注销.docx'))
    res = re.findall(r'〔(\d+)〕(\d+)', key)
    try:
        year = res[0][0]
        nb = res[0][1]
        company = value[0][1]
        com = value[0][0]
        month = value[0][3][1]
        day = value[0][3][2]
    except IndexError:
        continue
    replace_text(document, 'year', f'{year}')
    replace_text(document, 'nb', f'{nb}')
    replace_text(document, 'company', f'{company}')
    replace_text(document, 'month', f'{month}')
    replace_text(document, 'day', f'{day}')
    table = document.add_table(rows=len(value) + 1, cols=3, style="grid")  # 准备写入航班计划，新建表格
    table_para_style_input(table.cell(0, 0), '序号', font_pos='加急')
    table.cell(0, 0).width = Inches(0.2)
    table_para_style_input(table.cell(0, 1), '航线', font_pos='加急')
    table.cell(0, 1).width = Inches(9.5)
    table_para_style_input(table.cell(0, 2), '备注', font_pos='加急')
    table.cell(0, 2).width = Inches(0.5)
    for i, j in enumerate(value):
        air_line = j[2]
        table_para_style_input(table.cell(i + 1, 0), str(i + 1))
        table_para_style_input(table.cell(i + 1, 1), air_line)
    para_style_input(document, '注：1、此通知书一式两份；', font_pos='发往单位')
    para_style_input(document, '2、相关内容已在“中国民航航线航班信息管理系统”（http: // product.caachbjc.com）或WWW.CAAC.GOV.CN网站上予以公告。')
    title = f'{year}{nb}-{com}-注销'
    path_zhu = (os.path.join(PATH, '航线航班\航线经营许可\注销\{}'.format(year)))
    if not os.path.isdir(path_zhu):
        os.mkdir(path_zhu)
    path_zhu_file = os.path.join(path_zhu, '{}.docx'.format(title))
    if not os.path.isfile(path_zhu_file):
        document.save(path_zhu_file)
        print(path_zhu_file)

print('\n*********docx文件成功生成！用时{}秒......*********\n'.format(time.time() - time_start))
